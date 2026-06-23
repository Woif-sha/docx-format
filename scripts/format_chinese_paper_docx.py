#!/usr/bin/env python
"""Apply the user's default Chinese academic DOCX formatting rules."""

from __future__ import annotations

import argparse
import json
import platform
import re
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BODY_FONT_PT = 12
TITLE_FONT_PT = 18
FIRST_LEVEL_HEADING_FONT_PT = 14
CAPTION_FONT_PT = 10.5
TABLE_FONT_PT = 9
FOOTER_FONT_PT = 12
INDENT_PT = BODY_FONT_PT * 2
BLACK = RGBColor(0, 0, 0)

HEADING_RE = re.compile(r"^(?:[一二三四五六七八九十]+、|\d+(?:\.\d+)*\s+|\d+[.、])")
REF_RE = re.compile(r"^\[\d+\]")
CAPTION_RE = re.compile(r"^([图表])(?:\s*(\d+(?:[.-]\d+)?)\s*[:：.、-]?\s*(.*)|\s+(.+))$")
MAX_CAPTION_TITLE_CHARS = 32
CHINESE_NUMERAL = {
    "一": 1,
    "二": 2,
    "三": 3,
    "四": 4,
    "五": 5,
    "六": 6,
    "七": 7,
    "八": 8,
    "九": 9,
    "十": 10,
}


def has_drawing(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing")) or bool(paragraph._p.xpath(".//w:pict"))


def is_caption(text: str) -> bool:
    text = text.strip()
    return bool(CAPTION_RE.match(text))


def is_heading(text: str) -> bool:
    text = text.strip()
    return bool(text and HEADING_RE.match(text))


def is_first_level_heading(text: str) -> bool:
    text = text.strip()
    return bool(re.match(r"^[一二三四五六七八九十]+、", text))


def is_reference(text: str) -> bool:
    return bool(REF_RE.match(text.strip()))


def is_bibliography_lead(text: str) -> bool:
    text = text.strip()
    return "[J]." in text or "[R]." in text or "[EB/OL]" in text


def is_body(text: str, paragraph) -> bool:
    text = text.strip()
    if not text or has_drawing(paragraph):
        return False
    if is_caption(text) or is_heading(text) or is_reference(text) or is_bibliography_lead(text):
        return False
    return True


def iter_paragraphs(document: Document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def iter_block_items(document: Document):
    parent = document.element.body
    for child in parent.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def set_run_fonts(paragraph, size_pt: float = BODY_FONT_PT) -> None:
    for run in paragraph.runs:
        if run.text and all(char in "“”‘’" for char in run.text):
            set_quote_run_font(run)
        else:
            set_standard_run_font(run)
        run.font.size = Pt(size_pt)
        run.font.color.rgb = BLACK


def set_standard_run_font(run) -> None:
    r_fonts = get_or_add_r_fonts(run)
    run.font.name = "Times New Roman"
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:cs"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "宋体")


def set_quote_run_font(run) -> None:
    r_pr = run._element.get_or_add_rPr()
    r_fonts = get_or_add_r_fonts(run)
    run.font.name = "宋体"
    r_fonts.set(qn("w:ascii"), "宋体")
    r_fonts.set(qn("w:hAnsi"), "宋体")
    r_fonts.set(qn("w:cs"), "宋体")
    r_fonts.set(qn("w:eastAsia"), "宋体")
    r_fonts.set(qn("w:hint"), "eastAsia")
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:eastAsia"), "zh-CN")


def get_or_add_r_fonts(run):
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    return r_fonts


def clear_paragraph_content(paragraph) -> None:
    p_element = paragraph._p
    for child in list(p_element):
        if child.tag != qn("w:pPr"):
            p_element.remove(child)


def set_paragraph_plain_text(paragraph, text: str, size_pt: float = BODY_FONT_PT) -> None:
    clear_paragraph_content(paragraph)
    run = paragraph.add_run(text)
    set_standard_run_font(run)
    run.font.size = Pt(size_pt)
    run.font.color.rgb = BLACK


def set_single_line_spacing(paragraph) -> None:
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = 1.0


def set_one_point_five_spacing(paragraph) -> None:
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def add_page_number_field(paragraph) -> None:
    clear_paragraph_content(paragraph)
    run = paragraph.add_run()
    set_standard_run_font(run)
    run.font.size = Pt(FOOTER_FONT_PT)
    run.font.color.rgb = BLACK

    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    field_separate = OxmlElement("w:fldChar")
    field_separate.set(qn("w:fldCharType"), "separate")
    field_result = OxmlElement("w:t")
    field_result.text = "1"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")

    run._r.append(field_begin)
    run._r.append(instr_text)
    run._r.append(field_separate)
    run._r.append(field_result)
    run._r.append(field_end)


def apply_footer_page_number(section) -> None:
    section.footer_distance = Cm(1.27)
    section.different_first_page_header_footer = False
    footer = section.footer
    try:
        footer.is_linked_to_previous = False
    except ValueError:
        pass

    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    set_single_line_spacing(paragraph)
    add_page_number_field(paragraph)

    for extra in list(footer.paragraphs[1:]):
        extra._element.getparent().remove(extra._element)


def apply_document_title_format(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    set_one_point_five_spacing(paragraph)
    set_run_fonts(paragraph, size_pt=TITLE_FONT_PT)
    for run in paragraph.runs:
        run.bold = True


def rebuild_runs_with_quote_fonts(paragraph, text: str, size_pt: float = BODY_FONT_PT) -> None:
    clear_paragraph_content(paragraph)
    buffer = []
    for char in text:
        if char in "“”‘’":
            if buffer:
                run = paragraph.add_run("".join(buffer))
                set_standard_run_font(run)
                run.font.size = Pt(size_pt)
                buffer = []
            run = paragraph.add_run(char)
            set_quote_run_font(run)
            run.font.size = Pt(size_pt)
        else:
            buffer.append(char)
    if buffer:
        run = paragraph.add_run("".join(buffer))
        set_standard_run_font(run)
        run.font.size = Pt(size_pt)


def process_quotes(paragraph, size_pt: float = BODY_FONT_PT) -> int:
    text = paragraph.text
    converted = curly_quotes(text)
    if converted != text or any(char in converted for char in "“”‘’"):
        rebuild_runs_with_quote_fonts(paragraph, converted, size_pt=size_pt)
    return text.count('"')


def curly_quotes(text: str) -> str:
    out = []
    opening = True
    for char in text:
        if char in {'"', "＂"}:
            out.append("“" if opening else "”")
            opening = not opening
        else:
            out.append(char)
    return "".join(out)


def chinese_chapter_number(text: str, fallback: int) -> int:
    match = re.match(r"^([一二三四五六七八九十]+)、", text.strip())
    if not match:
        return fallback
    raw = match.group(1)
    if raw in CHINESE_NUMERAL:
        return CHINESE_NUMERAL[raw]
    if raw.startswith("十") and len(raw) == 2:
        return 10 + CHINESE_NUMERAL.get(raw[1], 0)
    if raw.endswith("十") and len(raw) == 2:
        return CHINESE_NUMERAL.get(raw[0], 1) * 10
    if "十" in raw:
        left, right = raw.split("十", 1)
        return CHINESE_NUMERAL.get(left, 1) * 10 + CHINESE_NUMERAL.get(right, 0)
    return fallback


def caption_kind(text: str) -> str | None:
    match = CAPTION_RE.match(text.strip())
    return match.group(1) if match else None


def extract_caption_title(text: str) -> str:
    match = CAPTION_RE.match(text.strip())
    if match:
        return (match.group(3) or match.group(4) or "").strip()
    return text.strip()


def normalize_caption_text(kind: str, chapter: int, index: int, text: str, fallback_title: str | None = None) -> str:
    title = sanitize_caption_title(extract_caption_title(text) or fallback_title or default_caption_title(kind, chapter))
    return f"{kind}{chapter}-{index}  {title}"


def default_caption_title(kind: str, chapter: int) -> str:
    return f"第{chapter}章相关{'图示' if kind == '图' else '数据'}"


def sanitize_caption_title(title: str) -> str:
    title = re.sub(r"\s+", "", title.strip())
    title = re.sub(r"^[：:，,。.、；;\-]+", "", title)
    title = re.sub(r"[。；;，,、：:]+$", "", title)
    if not title:
        return title
    return title[:MAX_CAPTION_TITLE_CHARS]


def clean_context_title(text: str) -> str:
    text = re.sub(r"^[一二三四五六七八九十]+、", "", text.strip())
    text = re.sub(r"^\d+(?:\.\d+)*[、.\s]*", "", text)
    text = re.sub(r"^\[\d+\]\s*", "", text)
    text = re.sub(r"\s+", "", text)
    text = re.split(r"[。；;！？!?]", text, maxsplit=1)[0]
    return sanitize_caption_title(text)


def table_header_title(table: Table) -> str:
    if not table.rows:
        return ""
    headers = []
    seen = set()
    for cell in table.rows[0].cells:
        text = clean_context_title(cell.text)
        if text and text not in seen:
            headers.append(text)
            seen.add(text)
    if not headers:
        return ""
    joined = "、".join(headers)
    return sanitize_caption_title(f"{joined}统计表")


def nearby_context_title(blocks: list, index: int, current_heading: str) -> str:
    for step in range(1, 4):
        previous_index = index - step
        if previous_index >= 0 and isinstance(blocks[previous_index], Paragraph):
            text = blocks[previous_index].text.strip()
            if text and not has_drawing(blocks[previous_index]) and not is_caption(text):
                candidate = clean_context_title(text)
                if candidate:
                    return candidate

    for step in range(1, 3):
        next_index = index + step
        if next_index < len(blocks) and isinstance(blocks[next_index], Paragraph):
            text = blocks[next_index].text.strip()
            if text and not has_drawing(blocks[next_index]) and not is_caption(text):
                candidate = clean_context_title(text)
                if candidate:
                    return candidate

    return clean_context_title(current_heading)


def proposed_caption_title(kind: str, blocks: list, index: int, current_heading: str, chapter: int) -> str:
    if kind == "表" and isinstance(blocks[index], Table):
        candidate = table_header_title(blocks[index])
        if candidate:
            return candidate
    candidate = nearby_context_title(blocks, index, current_heading)
    if candidate:
        suffix = "图示" if kind == "图" else "情况"
        if candidate.endswith(("图", "图示", "图例", "表", "情况", "数据", "统计表")):
            return candidate
        return sanitize_caption_title(f"{candidate}{suffix}")
    return default_caption_title(kind, chapter)


def insert_paragraph_before(element, text: str) -> Paragraph:
    paragraph_element = OxmlElement("w:p")
    anchor = getattr(element, "_element", None)
    if anchor is None:
        anchor = getattr(element, "_tbl", None)
    anchor.addprevious(paragraph_element)
    paragraph = Paragraph(paragraph_element, element._parent)
    set_paragraph_plain_text(paragraph, text, size_pt=CAPTION_FONT_PT)
    return paragraph


def insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    paragraph_element = OxmlElement("w:p")
    paragraph._p.addnext(paragraph_element)
    inserted = Paragraph(paragraph_element, paragraph._parent)
    set_paragraph_plain_text(inserted, text, size_pt=CAPTION_FONT_PT)
    return inserted


def ensure_numbered_captions(document: Document) -> dict:
    blocks = list(iter_block_items(document))
    current_chapter = 1
    current_heading = ""
    fallback_chapter = 1
    counters = {"图": 0, "表": 0}
    processed: set[int] = set()
    inserted = 0
    normalized = 0

    for index, block in enumerate(blocks):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if is_first_level_heading(text):
                current_chapter = chinese_chapter_number(text, fallback_chapter)
                current_heading = text
                fallback_chapter = current_chapter + 1
                counters = {"图": 0, "表": 0}
                continue

            if has_drawing(block):
                next_block = blocks[index + 1] if index + 1 < len(blocks) else None
                if isinstance(next_block, Paragraph) and caption_kind(next_block.text) == "图":
                    if id(next_block._p) not in processed:
                        counters["图"] += 1
                        set_paragraph_plain_text(
                            next_block,
                            normalize_caption_text("图", current_chapter, counters["图"], next_block.text),
                            size_pt=CAPTION_FONT_PT,
                        )
                        processed.add(id(next_block._p))
                        normalized += 1
                else:
                    counters["图"] += 1
                    title = proposed_caption_title("图", blocks, index, current_heading, current_chapter)
                    insert_paragraph_after(
                        block,
                        normalize_caption_text("图", current_chapter, counters["图"], title, fallback_title=title),
                    )
                    inserted += 1
                continue

            kind = caption_kind(text)
            if kind and id(block._p) not in processed:
                counters[kind] += 1
                set_paragraph_plain_text(
                    block,
                    normalize_caption_text(kind, current_chapter, counters[kind], text),
                    size_pt=CAPTION_FONT_PT,
                )
                processed.add(id(block._p))
                normalized += 1

        elif isinstance(block, Table):
            previous_block = blocks[index - 1] if index > 0 else None
            if isinstance(previous_block, Paragraph) and caption_kind(previous_block.text) == "表":
                if id(previous_block._p) not in processed:
                    counters["表"] += 1
                    set_paragraph_plain_text(
                        previous_block,
                        normalize_caption_text("表", current_chapter, counters["表"], previous_block.text),
                        size_pt=CAPTION_FONT_PT,
                    )
                    processed.add(id(previous_block._p))
                    normalized += 1
            else:
                counters["表"] += 1
                title = proposed_caption_title("表", blocks, index, current_heading, current_chapter)
                insert_paragraph_before(
                    block,
                    normalize_caption_text("表", current_chapter, counters["表"], title, fallback_title=title),
                )
                inserted += 1

    return {"inserted": inserted, "normalized": normalized}


def clear_shading(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for shd in list(tc_pr.findall(qn("w:shd"))):
        tc_pr.remove(shd)


def cell_border(cell, **borders) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, settings in borders.items():
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        for key, value in settings.items():
            element.set(qn(f"w:{key}"), str(value))


def apply_three_line_table(table) -> int:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    line = {"val": "single", "sz": "12", "space": "0", "color": "000000"}
    mid = {"val": "single", "sz": "6", "space": "0", "color": "000000"}
    nil = {"val": "nil"}
    quote_replacements = 0

    if not table.rows:
        return quote_replacements

    for row in table.rows:
        for cell in row.cells:
            clear_shading(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_border(cell, top=nil, left=nil, bottom=nil, right=nil, insideH=nil, insideV=nil)
            for paragraph in cell.paragraphs:
                quote_replacements += process_quotes(paragraph, size_pt=TABLE_FONT_PT)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.first_line_indent = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                set_single_line_spacing(paragraph)
                set_run_fonts(paragraph, size_pt=TABLE_FONT_PT)

    for cell in table.rows[0].cells:
        clear_shading(cell)
        cell_border(cell, top=line, bottom=mid, left=nil, right=nil)
    for cell in table.rows[-1].cells:
        cell_border(cell, bottom=line, left=nil, right=nil)
    return quote_replacements


def apply_word_com_fullwidth_quotes(path: Path) -> dict:
    """Use Word COM to mimic Find [“”] -> select all -> Aa -> full-width."""
    result = {"attempted": False, "ok": False, "matches": 0, "error": None}
    if platform.system() != "Windows":
        result["error"] = "not_windows"
        return result

    try:
        import win32com.client  # type: ignore
    except Exception as exc:  # pragma: no cover - depends on local Windows setup
        fallback = apply_word_com_fullwidth_quotes_powershell(path)
        fallback["pywin32_error"] = str(exc)
        return fallback

    result["attempted"] = True
    word = None
    doc = None
    try:
        word = win32com.client.gencache.EnsureDispatch("Word.Application")
        constants = win32com.client.constants
        wd_find_stop = getattr(constants, "wdFindStop", 0)
        wd_width_full_width = getattr(constants, "wdWidthFullWidth", 7)
        doc = word.Documents.Open(str(path.resolve()), ReadOnly=False, AddToRecentFiles=False)
        word.Visible = False
        rng = doc.Content
        find = rng.Find
        find.ClearFormatting()
        find.Text = "[“”]"
        find.Forward = True
        find.Wrap = wd_find_stop
        find.MatchWildcards = True

        while find.Execute():
            found = rng.Duplicate
            found.CharacterWidth = wd_width_full_width
            found.Font.Name = "宋体"
            found.Font.NameAscii = "宋体"
            found.Font.NameFarEast = "宋体"
            found.Font.NameOther = "宋体"
            result["matches"] += 1
            rng.Start = found.End
            rng.End = doc.Content.End

        doc.Save()
        result["ok"] = True
    except Exception as exc:  # pragma: no cover - depends on local Word setup
        result["error"] = str(exc)
    finally:
        if doc is not None:
            try:
                doc.Close(SaveChanges=True)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
    return result


def apply_word_com_fullwidth_quotes_powershell(path: Path) -> dict:
    """Fallback Word COM automation through PowerShell when pywin32 is unavailable."""
    result = {"attempted": True, "ok": False, "matches": 0, "error": None, "route": "powershell"}
    ps_script = r'''
param([string]$DocPath)
$ErrorActionPreference = "Stop"
$word = $null
$doc = $null
try {
  $word = New-Object -ComObject Word.Application
  $word.Visible = $false
  $doc = $word.Documents.Open($DocPath, $false, $false)
  $range = $doc.Content
  $find = $range.Find
  $find.ClearFormatting()
  $find.Text = '[“”]'
  $find.Forward = $true
  $find.Wrap = 0
  $find.MatchWildcards = $true
  $count = 0
  while ($find.Execute()) {
    $found = $range.Duplicate
    try { $found.CharacterWidth = 7 } catch {}
    $found.Font.Name = '宋体'
    $found.Font.NameAscii = '宋体'
    $found.Font.NameFarEast = '宋体'
    $found.Font.NameOther = '宋体'
    $count += 1
    $range.Start = $found.End
    $range.End = $doc.Content.End
  }
  $doc.Save()
  Write-Output "COUNT=$count"
}
finally {
  if ($doc -ne $null) { try { $doc.Close($true) } catch {} }
  if ($word -ne $null) { try { $word.Quit() } catch {} }
}
'''
    script_path = Path(tempfile.gettempdir()) / "chinese_paper_docx_quote_width.ps1"
    script_path.write_text(ps_script, encoding="utf-8-sig")
    try:
        completed = subprocess.run(
            [
                "powershell",
                "-NoProfile",
                "-ExecutionPolicy",
                "Bypass",
                "-File",
                str(script_path),
                "-DocPath",
                str(path.resolve()),
            ],
            text=True,
            capture_output=True,
            timeout=120,
        )
        if completed.returncode != 0:
            result["error"] = (completed.stderr or completed.stdout or "").strip()
            return result
        match = re.search(r"COUNT=(\d+)", completed.stdout)
        result["matches"] = int(match.group(1)) if match else 0
        result["ok"] = True
    except Exception as exc:  # pragma: no cover - depends on local Word setup
        result["error"] = str(exc)
    finally:
        try:
            script_path.unlink()
        except OSError:
            pass
    return result


def apply_document_format(path: Path, output: Path, use_word_com: bool = True) -> dict:
    document = Document(path)

    for section in document.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        apply_footer_page_number(section)

    title_paragraph_index = next(
        (index for index, paragraph in enumerate(document.paragraphs) if paragraph.text.strip() and not has_drawing(paragraph)),
        None,
    )

    quote_replacements = 0
    body_count = 0
    heading_count = 0
    reference_count = 0
    image_count = 0
    caption_count = 0
    caption_result = ensure_numbered_captions(document)

    for paragraph_index, paragraph in enumerate(document.paragraphs):
        stripped_before = paragraph.text.strip()
        if paragraph_index == title_paragraph_index:
            run_size = TITLE_FONT_PT
        elif is_caption(stripped_before):
            run_size = CAPTION_FONT_PT
        elif is_first_level_heading(stripped_before):
            run_size = FIRST_LEVEL_HEADING_FONT_PT
        else:
            run_size = BODY_FONT_PT

        quote_replacements += process_quotes(paragraph, size_pt=run_size)

        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        set_one_point_five_spacing(paragraph)

        stripped = paragraph.text.strip()
        if paragraph_index == title_paragraph_index:
            apply_document_title_format(paragraph)
        elif has_drawing(paragraph):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            set_run_fonts(paragraph, size_pt=BODY_FONT_PT)
            image_count += 1
        elif is_caption(stripped):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            set_single_line_spacing(paragraph)
            set_run_fonts(paragraph, size_pt=CAPTION_FONT_PT)
            caption_count += 1
        elif is_reference(stripped):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Pt(0)
            set_run_fonts(paragraph, size_pt=BODY_FONT_PT)
            reference_count += 1
        elif is_heading(stripped):
            paragraph.alignment = paragraph.alignment or WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            heading_size = FIRST_LEVEL_HEADING_FONT_PT if is_first_level_heading(stripped) else BODY_FONT_PT
            set_run_fonts(paragraph, size_pt=heading_size)
            heading_count += 1
        elif is_body(stripped, paragraph):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Pt(INDENT_PT)
            set_run_fonts(paragraph, size_pt=BODY_FONT_PT)
            body_count += 1
        else:
            set_run_fonts(paragraph, size_pt=BODY_FONT_PT)

    for table in document.tables:
        quote_replacements += apply_three_line_table(table)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    word_com_result = apply_word_com_fullwidth_quotes(output) if use_word_com else {
        "attempted": False,
        "ok": False,
        "matches": 0,
        "error": "disabled",
    }

    return {
        "output": str(output),
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "body_paragraphs": body_count,
        "headings": heading_count,
        "references": reference_count,
        "captions": caption_count,
        "image_paragraphs": image_count,
        "captions_inserted": caption_result["inserted"],
        "captions_normalized": caption_result["normalized"],
        "quote_replacements": quote_replacements,
        "word_com_fullwidth_quotes": word_com_result,
    }


def validate_docx(path: Path) -> dict:
    result = {
        "zip_ok": zipfile.is_zipfile(path),
        "ascii_double_quotes": None,
        "paragraphs": None,
        "tables": None,
    }
    if not result["zip_ok"]:
        return result
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in iter_paragraphs(document))
    result.update(
        {
            "ascii_double_quotes": text.count('"'),
            "paragraphs": len(document.paragraphs),
            "tables": len(document.tables),
        }
    )
    return result


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path)
    parser.add_argument("--output", type=Path)
    parser.add_argument("--json", action="store_true")
    parser.add_argument(
        "--skip-word-com-fullwidth-quotes",
        action="store_true",
        help="Do not attempt the Windows Word COM full-width quote pass.",
    )
    args = parser.parse_args()

    output = args.output or args.input
    summary = apply_document_format(args.input, output, use_word_com=not args.skip_word_com_fullwidth_quotes)
    summary["validation"] = validate_docx(output)

    if args.json:
        print(json.dumps(summary, ensure_ascii=False, indent=2))
    else:
        print(f"Formatted: {output}")
        print(f"Paragraphs: {summary['paragraphs']}; tables: {summary['tables']}")
        print(f"ASCII double quotes after formatting: {summary['validation']['ascii_double_quotes']}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
